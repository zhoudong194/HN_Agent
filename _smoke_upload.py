"""Smoke-test the /api/ingest upload endpoint."""
import io
import urllib.request
import urllib.error
import json
from docx import Document

# Build a small docx in memory
doc = Document()
doc.add_heading("测试新规", level=1)
doc.add_paragraph("员工每年享受 12 天带薪事假，用于处理家庭事务。")
doc.add_paragraph("申请流程：提前 3 个工作日报直属上级。")
buf = io.BytesIO()
doc.save(buf)
buf.seek(0)

boundary = "----testboundary1234"
body = (
    f"--{boundary}\r\n"
    "Content-Disposition: form-data; name=\"file\"; filename=\"test_policy.docx\"\r\n"
    "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n"
    "\r\n"
).encode("utf-8") + buf.read() + f"\r\n--{boundary}--\r\n".encode("utf-8")

req = urllib.request.Request(
    "http://127.0.0.1:8000/api/ingest",
    data=body,
    headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
    method="POST",
)

try:
    with urllib.request.urlopen(req, timeout=120) as r:
        print(f"Status: {r.status}")
        out = json.loads(r.read())
        print(json.dumps(out, ensure_ascii=False, indent=2))
except urllib.error.HTTPError as e:
    print(f"HTTPError: {e.code}")
    print(e.read().decode("utf-8"))
