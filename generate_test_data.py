"""
generate_test_data.py - Generate test .docx file with company rules
Run this script to create sample company policy documents in ./data directory.
"""

import sys
import io

# Fix Windows console encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_company_rules_docx():
    """Create a sample company rules document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('员工手册 - 公司规章制度', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document info
    doc.add_paragraph('版本：V2.0')
    doc.add_paragraph('生效日期：2024年1月1日')
    doc.add_paragraph()
    
    # Chapter 1: Leave Policy
    doc.add_heading('第一章 年假制度', level=1)
    
    doc.add_heading('1.1 年假天数', level=2)
    doc.add_paragraph('公司员工享受带薪年假，具体天数如下：')
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Table data
    headers = table.rows[0].cells
    headers[0].text = '工龄'
    headers[1].text = '年假天数'
    
    data = [
        ('1年以下（含1年）', '5天'),
        ('1年以上-3年以下', '7天'),
        ('3年以上-5年以下', '10天'),
        ('5年以上', '15天'),
    ]
    
    for i, (work_years, days) in enumerate(data, start=1):
        row = table.rows[i].cells
        row[0].text = work_years
        row[1].text = days
    
    doc.add_paragraph()
    
    doc.add_heading('1.2 年假申请流程', level=2)
    doc.add_paragraph('员工申请年假需遵循以下流程：')
    doc.add_paragraph('1. 提前5个工作日在OA系统提交申请')
    doc.add_paragraph('2. 由直属上级审批')
    doc.add_paragraph('3. 部门经理最终确认')
    doc.add_paragraph('4. HR系统登记备案')
    
    doc.add_heading('1.3 年假使用规定', level=2)
    doc.add_paragraph('• 年假应在当年度内使用，原则上不跨年累积')
    doc.add_paragraph('• 因工作原因无法休完年假的，经批准后可延期至次年3月31日前使用')
    doc.add_paragraph('• 离职时未休年假按日薪的3倍折算补偿')
    
    # Chapter 2: Attendance
    doc.add_heading('第二章 考勤制度', level=1)
    
    doc.add_heading('2.1 工作时间', level=2)
    doc.add_paragraph('公司实行标准工时制，具体工作时间如下：')
    doc.add_paragraph('• 周一至周五：上午9:00 - 12:00，下午13:30 - 18:00')
    doc.add_paragraph('• 午休时间：12:00 - 13:30（不计入工作时间）')
    doc.add_paragraph('• 周六、周日为休息日')
    
    doc.add_heading('2.2 迟到早退处理', level=2)
    doc.add_paragraph('员工上下班必须打卡，具体规定：')
    doc.add_paragraph('• 迟到30分钟以内：扣罚当日餐补')
    doc.add_paragraph('• 迟到30分钟-2小时：按半天事假处理')
    doc.add_paragraph('• 迟到2小时以上：按旷工半天处理')
    doc.add_paragraph('• 早退30分钟以上：按旷工半天处理')
    
    doc.add_heading('2.3 加班规定', level=2)
    doc.add_paragraph('公司不鼓励加班，如因工作需要加班，需遵守以下规定：')
    doc.add_paragraph('• 工作日加班：按小时支付1.5倍工资')
    doc.add_paragraph('• 周末加班：按小时支付2倍工资')
    doc.add_paragraph('• 法定节假日加班：按小时支付3倍工资')
    doc.add_paragraph('• 加班需提前填写加班申请单，经审批后生效')
    
    # Chapter 3: Expenses
    doc.add_heading('第三章 报销制度', level=1)
    
    doc.add_heading('3.1 差旅报销', level=2)
    doc.add_paragraph('员工因公出差可报销以下费用：')
    doc.add_paragraph('• 交通费：飞机票（经济舱）、火车票（高铁二等座）')
    doc.add_paragraph('• 住宿费：高管每日不超过800元，普通员工每日不超过400元')
    doc.add_paragraph('• 餐饮费：高管每日不超过200元，普通员工每日不超过100元')
    doc.add_paragraph('• 其他合理费用：需提供发票并说明用途')
    
    doc.add_heading('3.2 报销流程', level=2)
    doc.add_paragraph('1. 在费用发生后15日内提交报销申请')
    doc.add_paragraph('2. 附上正规发票和费用明细')
    doc.add_paragraph('3. 直属上级审核')
    doc.add_paragraph('4. 财务部门复核')
    doc.add_paragraph('5. 总经理审批（单笔超过5000元需此步骤）')
    doc.add_paragraph('6. 出纳付款（通常在审批后3个工作日内）')
    
    # Chapter 4: Office Supplies
    doc.add_heading('第四章 办公用品管理', level=1)
    
    doc.add_heading('4.1 办公用品申领', level=2)
    doc.add_paragraph('员工可通过OA系统申领办公用品：')
    doc.add_paragraph('• 常规用品（如笔、本子、文件夹等）：每月限领一次')
    doc.add_paragraph('• 电子设备（如鼠标、键盘等）：需部门经理特批')
    doc.add_paragraph('• 打印耗材：按实际需求申领，需登记用途')
    
    doc.add_heading('4.2 办公设备管理', level=2)
    doc.add_paragraph('• 公司配置的办公设备需妥善保管')
    doc.add_paragraph('• 设备损坏或丢失需及时报告IT部门')
    doc.add_paragraph('• 离职时需归还所有公司设备')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('本手册由人力资源部负责解释，如有疑问请联系HR部门。')
    doc.add_paragraph('联系邮箱：hr@company.com')
    
    return doc


def main():
    # Ensure data directory exists
    os.makedirs('./data', exist_ok=True)
    
    # Create document
    doc = create_company_rules_docx()
    
    # Save
    output_path = './data/员工手册_公司规章制度.docx'
    doc.save(output_path)
    print(f"[OK] Test document created: {output_path}")
    
    # Also create a simple PDF text file for testing
    pdf_content = """公司采购管理制度

第一章 总则
第一条 为规范公司采购行为，降低采购成本，提高经济效益，特制定本制度。

第二条 适用范围
本制度适用于公司所有部门及子公司的物资采购活动。

第三条 采购原则
1. 公开、公平、公正原则
2. 质优价廉原则
3. 性价比最优原则

第二章 采购流程
第四条 采购申请
- 使用部门提出采购申请，填写《采购申请表》
- 采购金额超过5000元需部门经理审批
- 采购金额超过20000元需总经理审批

第五条 供应商选择
- 采购部门负责供应商的筛选和评估
- 原则上每类物资需保持2-3家合格供应商
- 定期对供应商进行绩效考核

第六条 采购执行
- 合同金额小于10000元：采购员可直接执行
- 合同金额10000-50000元：需签订采购合同
- 合同金额超过50000元：需进行招标程序

第三章 付款方式
第七条 付款期限
- 月结30天：适用于长期供应商
- 货到付款：适用于小额采购
- 预付款：仅适用于设备采购，需特批

第四章 附则
第八条 本制度由财务部负责解释。
第九条 本制度自发布之日起施行。
"""
    
    # Save as markdown (simulating converted PDF content)
    pdf_path = './data/采购管理制度.md'
    with open(pdf_path, 'w', encoding='utf-8') as f:
        f.write(pdf_content)
    print(f"[OK] Test markdown file created: {pdf_path}")
    
    print("\n[*] Test data files created successfully!")
    print("   - ./data/员工手册_公司规章制度.docx")
    print("   - ./data/采购管理制度.md")


if __name__ == "__main__":
    main()
