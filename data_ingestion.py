"""
data_ingestion.py - Offline ingestion script for company rules RAG system.

v4 使用 PostgreSQL + pgvector 进行向量检索和存储。
HNSW 索引在 INSERT 时自动维护，无需重建步骤。

Pipeline:
    文件扫描 → 文本提取 → 语义切块 → BGE 向量化 → PostgreSQL 入库
"""

from __future__ import annotations

import re

import hashlib
import io
import os
import sys
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import List

# Fix Windows console encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

import config

# Document parsing
import docx
import docx2txt
try:
    import textract
except ImportError:  # optional; old .doc parsing is best-effort
    textract = None
from llama_index.core import Document
# Embedding model
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# Database (PostgreSQL + pgvector)
import database

# Configuration
DATA_DIR = config.DATA_DIR
EMBED_MODEL_NAME = config.EMBED_MODEL_NAME
EMBED_DIM = config.EMBED_DIM


@dataclass
class StructuredChunk:
    text: str
    metadata: dict

    def get_text(self) -> str:
        return self.text


def doc_to_text(file_path: str) -> str:
    """Extract text from legacy .doc file. Uses textract first, falls back to docx2txt."""
    if textract is not None:
        try:
            text = textract.process(file_path, extension="doc", encoding="utf-8")
            decoded = text.decode("utf-8", errors="replace")
            q_ratio = decoded.count("?") / max(len(decoded), 1)
            if q_ratio > 0.1:
                raise ValueError(f"High question-mark ratio: {q_ratio:.1%}")
            return decoded
        except Exception:
            pass

    # Fallback: try docx2txt on the file. This only works for OOXML-like files.
    try:
        text = docx2txt.process(file_path)
        return text or ""
    except Exception:
        return ""


def docx_to_text(file_path: str) -> str:
    """Extract text from DOCX file using python-docx, preserving headings."""
    doc = docx.Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            level = para.style.name.split(" ")[-1]
            try:
                level = int(level)
                text_parts.append(f"{'#' * (level + 1)} {para.text}")
            except ValueError:
                text_parts.append(f"## {para.text}")
        else:
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(f"| {row_text} |")
        text_parts.append("")

    return "\n".join(text_parts)


def convert_to_markdown(file_path: str) -> List[Document]:
    """Convert DOCX/MD files to Markdown format documents."""
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".docx":
            text_content = docx_to_text(str(file_path))
            text_content = f"# {file_path.stem}\n\n{text_content}"
            doc = Document(text=text_content)
            doc.metadata["source_file"] = str(file_path)
            doc.metadata["file_type"] = suffix
            return [doc]

        elif suffix == ".doc":
            text_content = doc_to_text(str(file_path))
            if not text_content.strip():
                raise ValueError("textract returned empty text")
            text_content = f"# {file_path.stem}\n\n{text_content}"
            doc = Document(text=text_content)
            doc.metadata["source_file"] = str(file_path)
            doc.metadata["file_type"] = suffix
            return [doc]

        elif suffix == ".md":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            doc = Document(text=content)
            doc.metadata["source_file"] = str(file_path)
            doc.metadata["file_type"] = suffix
            return [doc]

        else:
            print(f"Skipping unsupported file type: {file_path}")
            return []

    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return []


def scan_data_directory(data_dir: str) -> List[tuple[Path, Document]]:
    """Scan data directory for supported files. Returns (Path, Document) pairs."""
    data_path = Path(data_dir)
    results: List[tuple[Path, Document]] = []
    seen: set = set()
    supported = [".docx", ".doc", ".pdf", ".md"]

    for ext in supported:
        for file_path in data_path.glob(f"*{ext}"):
            if file_path in seen:
                continue
            seen.add(file_path)
            print(f"Processing: {file_path.name}")
            docs = convert_to_markdown(str(file_path))
            for doc in docs:
                results.append((file_path, doc))

    return results


def _split_long_text(text: str, max_chars: int = 700, overlap: int = 100) -> List[str]:
    text = text.strip()
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]

    chunks: List[str] = []
    step = max(1, max_chars - overlap)
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start += step
    return chunks


def create_semantic_chunks(documents: List[Document]) -> List[StructuredChunk]:
    """
    Structure-first chunking for policy documents.

    Preferred hierarchy:
      title → chapter → article → paragraph/list/table
    """
    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    chunks: List[StructuredChunk] = []

    for doc in documents:
        source_file = doc.metadata.get("source_file")
        file_stem = Path(source_file).stem if source_file else "document"
        raw_text = getattr(doc, "text", None) or doc.get_content()
        lines = raw_text.splitlines()

        heading_stack = [file_stem, None, None]
        section_lines: List[str] = []
        chunk_index = 0

        def flush_section():
            nonlocal chunk_index
            body = "\n".join(line for line in section_lines if line is not None).strip()
            if not body:
                return

            header_1 = heading_stack[0] or file_stem
            header_2 = heading_stack[1]
            header_3 = heading_stack[2]
            for part in _split_long_text(body):
                chunks.append(
                    StructuredChunk(
                        text=part,
                        metadata={
                            "source_file": source_file,
                            "file_type": doc.metadata.get("file_type"),
                            "header_1": header_1,
                            "header_2": header_2,
                            "header_3": header_3,
                            "chunk_index": chunk_index,
                        },
                    )
                )
                chunk_index += 1

        for line in lines:
            match = heading_re.match(line.strip())
            if match:
                flush_section()
                hashes, title = match.groups()
                level = min(len(hashes), 6)
                if level == 1:
                    heading_stack[0] = title.strip()
                    heading_stack[1] = None
                    heading_stack[2] = None
                elif level == 2:
                    heading_stack[1] = title.strip()
                    heading_stack[2] = None
                elif level == 3:
                    heading_stack[2] = title.strip()
                else:
                    heading_stack[2] = title.strip()
                section_lines = []
                continue

            if line.strip():
                section_lines.append(line.rstrip())
            elif section_lines and section_lines[-1] != "":
                section_lines.append("")

        flush_section()

    return chunks


def setup_embedding_model() -> HuggingFaceEmbedding:
    """Initialize BGE embedding model for Chinese text."""
    embed_model = HuggingFaceEmbedding(
        model_name=EMBED_MODEL_NAME,
        query_instruction="为这个句子生成表示以用于检索相关文章：",
        text_instruction="把这段文章转化成一个向量表示：",
        embed_batch_size=32,
    )
    return embed_model


def ingest_to_database():
    """Main ingestion pipeline: scan → chunk → embed → PostgreSQL / pgvector."""
    print("=" * 60)
    print("Starting RAG Data Ingestion Pipeline (PostgreSQL + pgvector)")
    print("=" * 60)

    # Step 1: Embedding model
    print("\n[1/5] Initializing BGE embedding model...")
    embed_model = setup_embedding_model()
    print(f"  ✓ Model: {EMBED_MODEL_NAME} (dim={EMBED_DIM})")

    # Step 2: Scan documents
    print(f"\n[2/5] Scanning {DATA_DIR} for documents...")
    file_doc_pairs = scan_data_directory(DATA_DIR)
    if not file_doc_pairs:
        print("No documents found. Please add .docx / .md files to the data directory.")
        sys.exit(1)
    print(f"  ✓ Found {len(file_doc_pairs)} files")

    # Step 3: Structure-first chunking
    print("\n[3/5] Creating semantic chunks...")
    documents = [doc for _, doc in file_doc_pairs]
    nodes = create_semantic_chunks(documents)
    print(f"  ✓ Created {len(nodes)} semantic chunks")

    # Step 4: Embed + insert into FAISS + SQLite
    print(f"\n[4/5] Embedding & inserting {len(nodes)} chunks...")

    for file_path, _ in file_doc_pairs:
        content = file_path.read_bytes()
        doc_record, is_new = database.create_document(
            filename=file_path.name,
            file_type=file_path.suffix.lower(),
            file_size=len(content),
            content=content,
            category=None,
            uploader="system",
            title=file_path.stem,
        )

        if not is_new:
            print(f"  = [already exists] {file_path.name}")
            continue

        doc_id = doc_record["id"]
        print(f"  + [new] {file_path.name} (id={doc_id[:8]}...)")

        # Collect all nodes for this document
        doc_nodes = [n for n in nodes if n.metadata.get("source_file") == str(file_path)]

        chunk_records = []
        skipped = 0
        for i, node in enumerate(doc_nodes):
            text = node.get_text()
            if not text:
                skipped += 1
                continue

            # Content quality gate
            stripped = text.strip()
            if len(stripped) < 30:
                skipped += 1
                continue
            cjk = len(re.findall(r"[\u4e00-\u9fff]", stripped))
            if cjk / max(len(stripped), 1) < 0.15:
                skipped += 1
                continue

            vec = embed_model.get_text_embedding(text)
            chunk_records.append({
                "document_id": doc_id,
                "text": text,
                "vec": vec,
                "header_1": node.metadata.get("header_1"),
                "header_2": node.metadata.get("header_2"),
                "header_3": node.metadata.get("header_3"),
                "source_file": str(file_path),
                "chunk_index": i,
            })

        if chunk_records:
            n = database.insert_chunks_batch(chunk_records)
            print(f"      -> inserted {n} chunks" + (f" (skipped {skipped} low-quality)" if skipped else ""))

    # Rebuild BM25 index after ingestion
    print("\n[5/5] Building BM25 index...")
    import recall
    recall.rebuild_bm25()
    print("  ✓ BM25 index built")

    # Summary
    total_chunks = database.get_chunk_count()
    print(f"  ✓ {total_chunks} chunks indexed in PostgreSQL")
    print(f"  ✓ pgvector HNSW index auto-maintained")
    print(f"  ✓ BM25 sparse index in memory")

    print("\n" + "=" * 60)
    print(f"[OK] Ingestion Complete!")
    print(f"   - Files processed: {len(file_doc_pairs)}")
    print(f"   - Total chunks indexed: {total_chunks}")
    print("=" * 60)


if __name__ == "__main__":
    ingest_to_database()
