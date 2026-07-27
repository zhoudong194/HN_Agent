"""
data_ingestion.py - Offline ingestion script for company rules RAG system.

v3 使用 FAISS IndexHNSWFlat 进行向量检索（替代 ChromaDB / pgvector）。
元数据存储在 SQLite，无需额外部署数据库服务。

Pipeline:
    文件扫描 → 文本提取 → 语义切块 → BGE 向量化 → FAISS 索引 + SQLite 入库
"""

from __future__ import annotations

import hashlib
import io
import os
import sys
import uuid
from pathlib import Path
from typing import List

# Fix Windows console encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

import config

# Document parsing
import docx
import textract
from llama_index.core import Document
from llama_index.core.node_parser import MarkdownNodeParser
from llama_index.core.schema import MetadataMode

# Embedding model
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# Database (FAISS + SQLite)
import database

# Configuration
DATA_DIR = config.DATA_DIR
EMBED_MODEL_NAME = config.EMBED_MODEL_NAME
EMBED_DIM = config.EMBED_DIM


def doc_to_text(file_path: str) -> str:
    """Extract text from legacy .doc file using textract."""
    try:
        text = textract.process(file_path, extension="doc", encoding="utf-8")
        return text.decode("utf-8", errors="replace")
    except Exception as e:
        print(f"  [!] textract failed for {file_path}: {e}")
        return ""


def docx_to_text(file_path: str) -> str:
    """Extract text from DOCX file using python-docx, preserving headings."""
    doc = docx.Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            level = para.style.name.split(" ")[-1]
            try:
                level = int(level)
                text_parts.append(f"{'#' * (level + 1)} {para.text}")
            except ValueError:
                text_parts.append(f"## {para.text}")
        else:
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(f"| {row_text} |")
        text_parts.append("")

    return "\n".join(text_parts)


def convert_to_markdown(file_path: str) -> List[Document]:
    """Convert DOCX/MD files to Markdown format documents."""
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".docx":
            text_content = docx_to_text(str(file_path))
            text_content = f"# {file_path.stem}\n\n{text_content}"
            doc = Document(text=text_content)
            doc.metadata["source_file"] = str(file_path)
            doc.metadata["file_type"] = suffix
            return [doc]

        elif suffix == ".doc":
            text_content = doc_to_text(str(file_path))
            if not text_content.strip():
                raise ValueError("textract returned empty text")
            text_content = f"# {file_path.stem}\n\n{text_content}"
            doc = Document(text=text_content)
            doc.metadata["source_file"] = str(file_path)
            doc.metadata["file_type"] = suffix
            return [doc]

        elif suffix == ".md":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            doc = Document(text=content)
            doc.metadata["source_file"] = str(file_path)
            doc.metadata["file_type"] = suffix
            return [doc]

        else:
            print(f"Skipping unsupported file type: {file_path}")
            return []

    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return []


def scan_data_directory(data_dir: str) -> List[tuple[Path, Document]]:
    """Scan data directory for supported files. Returns (Path, Document) pairs."""
    data_path = Path(data_dir)
    results: List[tuple[Path, Document]] = []
    seen: set = set()
    supported = [".docx", ".doc", ".pdf", ".md"]

    for ext in supported:
        for file_path in data_path.glob(f"*{ext}"):
            if file_path in seen:
                continue
            seen.add(file_path)
            print(f"Processing: {file_path.name}")
            docs = convert_to_markdown(str(file_path))
            for doc in docs:
                results.append((file_path, doc))

    return results


def create_semantic_chunks(documents: List[Document]) -> List:
    """Parse documents into semantic chunks using MarkdownNodeParser."""
    parser = MarkdownNodeParser(
        include_metadata=True,
        include_prev_next_rel=True,
        metadata_mode=MetadataMode.ALL,
    )
    nodes = parser.get_nodes_from_documents(documents)
    return nodes


def setup_embedding_model() -> HuggingFaceEmbedding:
    """Initialize BGE embedding model for Chinese text."""
    embed_model = HuggingFaceEmbedding(
        model_name=EMBED_MODEL_NAME,
        query_instruction="为这个句子生成表示以用于检索相关文章：",
        text_instruction="把这段文章转化成一个向量表示：",
        embed_batch_size=32,
    )
    return embed_model


def ingest_to_database():
    """Main ingestion pipeline: scan → chunk → embed → FAISS + SQLite."""
    print("=" * 60)
    print("Starting RAG Data Ingestion Pipeline (FAISS + SQLite)")
    print("=" * 60)

    # Step 1: Embedding model
    print("\n[1/5] Initializing BGE embedding model...")
    embed_model = setup_embedding_model()
    print(f"  ✓ Model: {EMBED_MODEL_NAME} (dim={EMBED_DIM})")

    # Step 2: Scan documents
    print(f"\n[2/5] Scanning {DATA_DIR} for documents...")
    file_doc_pairs = scan_data_directory(DATA_DIR)
    if not file_doc_pairs:
        print("No documents found. Please add .docx / .md files to the data directory.")
        sys.exit(1)
    print(f"  ✓ Found {len(file_doc_pairs)} files")

    # Step 3: Semantic chunking
    print("\n[3/5] Creating semantic chunks...")
    documents = [doc for _, doc in file_doc_pairs]
    nodes = create_semantic_chunks(documents)
    print(f"  ✓ Created {len(nodes)} semantic chunks")

    # Step 4: Embed + insert into FAISS + SQLite
    print(f"\n[4/5] Embedding & inserting {len(nodes)} chunks...")

    for file_path, _ in file_doc_pairs:
        content = file_path.read_bytes()
        doc_record, is_new = database.create_document(
            filename=file_path.name,
            file_type=file_path.suffix.lower(),
            file_size=len(content),
            content=content,
            category=None,
            uploader="system",
            title=file_path.stem,
        )

        if not is_new:
            print(f"  = [already exists] {file_path.name}")
            continue

        doc_id = doc_record["id"]
        print(f"  + [new] {file_path.name} (id={doc_id[:8]}...)")

        # Collect all nodes for this document
        doc_nodes = [
            n for n in nodes
            if n.metadata.get("source_file") == str(file_path)
        ]

        chunk_records = []
        for i, node in enumerate(doc_nodes):
            text = node.get_text()
            if not text or len(text.strip()) < 10:
                continue

            vec = embed_model.get_text_embedding(text)
            chunk_records.append({
                "document_id": doc_id,
                "text": text,
                "vec": vec,
                "header_1": node.metadata.get("header_1"),
                "header_2": node.metadata.get("header_2"),
                "header_3": node.metadata.get("header_3"),
                "source_file": str(file_path),
                "chunk_index": i,
            })

        if chunk_records:
            n = database.insert_chunks_batch(chunk_records)
            print(f"      → inserted {n} chunks")

    # Step 5: Rebuild FAISS index
    print("\n[5/5] Building FAISS HNSW index...")
    database.rebuild_index()

    # Summary
    total_chunks = database.get_chunk_count()
    print(f"  ✓ FAISS index rebuilt with {total_chunks} total chunks")
    print(f"  ✓ SQLite metadata: {database.META_DB_PATH}")
    print(f"  ✓ FAISS index: {database.INDEX_FILE}")

    print("\n" + "=" * 60)
    print(f"[OK] Ingestion Complete!")
    print(f"   - Files processed: {len(file_doc_pairs)}")
    print(f"   - Total chunks indexed: {total_chunks}")
    print("=" * 60)


if __name__ == "__main__":
    ingest_to_database()
